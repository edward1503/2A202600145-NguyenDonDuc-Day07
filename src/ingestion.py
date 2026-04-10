from __future__ import annotations

import hashlib
from io import BytesIO
from pathlib import Path
from typing import Any

from .chunking import FixedSizeChunker, RecursiveChunker, SentenceChunker
from .models import Document


def load_documents_from_paths(file_paths: list[str]) -> list[Document]:
    """Load .txt/.md/.pdf files from disk into Document objects."""
    documents: list[Document] = []

    for raw_path in file_paths:
        path = Path(raw_path)
        if not path.exists() or not path.is_file():
            continue
        documents.extend(
            _documents_from_bytes(
                file_name=path.name,
                payload=path.read_bytes(),
                source=str(path),
                origin="local_path",
            )
        )

    return documents


def load_documents_from_uploads(uploaded_files: list[Any]) -> list[Document]:
    """Load uploaded .txt/.md/.pdf files into Document objects."""
    documents: list[Document] = []
    for file_obj in uploaded_files:
        filename = getattr(file_obj, "name", "")
        if not filename:
            continue
        payload = file_obj.getvalue()
        documents.extend(
            _documents_from_bytes(
                file_name=filename,
                payload=payload,
                source=f"uploaded/{filename}",
                origin="upload",
            )
        )
    return documents


def document_checksum(content: str) -> str:
    return hashlib.sha256(content.encode("utf-8")).hexdigest()


def document_signature(doc: Document) -> str:
    source = str((doc.metadata or {}).get("source", doc.id))
    return f"{source}:{document_checksum(doc.content)}"


def _documents_from_bytes(file_name: str, payload: bytes, source: str, origin: str) -> list[Document]:
    suffix = Path(file_name).suffix.lower()
    if suffix not in {".txt", ".md", ".pdf", ".docx"}:
        return []

    content = ""
    parser_used = "none"

    if suffix == ".pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=payload, filetype="pdf")
            content = "\n".join(page.get_text() for page in doc)
            parser_used = "pymupdf"
        except Exception as e:
            content = f"PDF extraction failed: {e}"
            parser_used = "failed"
    elif suffix == ".docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(BytesIO(payload))
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            parser_used = "python-docx"
        except Exception as e:
            content = f"DOCX extraction failed: {e}"
            parser_used = "failed"
    elif suffix in {".txt", ".md"}:
        content = payload.decode("utf-8", errors="ignore")
        parser_used = "utf-8"

    if not content.strip():
        return []

    metadata = {
        "source": source,
        "extension": suffix,
        "origin": origin,
        "parser_used": parser_used
    }
    
    metadata["checksum"] = document_checksum(content)
    return [
        Document(
            id=Path(file_name).stem,
            content=content,
            metadata=metadata,
        )
    ]


def create_chunker(method: str, params: dict[str, Any]):
    """Factory for supported chunking methods."""
    method = method.strip().lower()
    if method == "fixed_size":
        return FixedSizeChunker(
            chunk_size=int(params.get("chunk_size", 500)),
            overlap=int(params.get("overlap", 50)),
        )
    if method == "by_sentences":
        return SentenceChunker(
            max_sentences_per_chunk=int(params.get("max_sentences_per_chunk", 3)),
        )
    if method == "recursive":
        return RecursiveChunker(
            chunk_size=int(params.get("chunk_size", 500)),
            separators=params.get("separators"),
        )
    raise ValueError(f"Unsupported chunking method: {method}")


def chunk_documents(
    docs: list[Document],
    method: str,
    params: dict[str, Any] | None = None,
) -> list[Document]:
    """
    Convert full documents into chunk-level documents for indexing.

    Metadata added for traceability:
        - doc_id: original document id
        - chunk_index: index within original document
        - chunk_method: selected chunking method
    """
    params = params or {}
    chunker = create_chunker(method=method, params=params)
    out: list[Document] = []

    for doc in docs:
        chunks = chunker.chunk(doc.content)
        if not chunks:
            continue
        for idx, chunk_text in enumerate(chunks):
            chunk_metadata = dict(doc.metadata or {})
            chunk_metadata["doc_id"] = doc.id
            chunk_metadata["chunk_index"] = idx
            chunk_metadata["chunk_method"] = method
            out.append(
                Document(
                    id=f"{doc.id}::chunk_{idx}",
                    content=chunk_text,
                    metadata=chunk_metadata,
                )
            )

    return out
